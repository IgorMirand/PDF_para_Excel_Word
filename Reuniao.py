import ctypes
import sys
import os
import datetime
import re

import camelot
import pandas as pd
import openpyxl
import pdfplumber
import tempfile
import traceback

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
from openpyxl.styles import Font, PatternFill, Alignment

from PyQt5.QtCore import QTimer, QThread, pyqtSignal, Qt
from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QLabel, QPushButton,
    QFileDialog, QRadioButton, QHBoxLayout, QProgressBar,
    QGroupBox, QMessageBox, QLineEdit
)
from PyQt5.QtGui import QFont, QIcon
from openpyxl.utils import get_column_letter

# -----------------------
# Helper: global excepthook to log uncaught exceptions
# -----------------------
'''
def global_excepthook(exc_type, exc_value, exc_tb):
    err = ''.join(traceback.format_exception(exc_type, exc_value, exc_tb))
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    try:
        with open("app_log.txt", "a", encoding="utf-8") as lf:
            lf.write(f"[{timestamp}] UNCAUGHT EXCEPTION:\n{err}\n")
    except Exception:
        print("Could not write to log file.")
    # Print to console as well
    print(err)

sys.excepthook = global_excepthook
'''
# -----------------------
# Dates
# -----------------------
data = datetime.date.today().strftime("%d_%m_%Y")
data1 = datetime.date.today().strftime("%d/%m/%Y")

# -----------------------
# PDF Extractor
# -----------------------
class PDFExtractor:
    def __init__(self, pdf_file_object):
        self.pdf_file = pdf_file_object

    def _validate_pdf_magic_number(self):
        try:
            initial_pos = self.pdf_file.tell()
            self.pdf_file.seek(0)
            magic_number = self.pdf_file.read(5)
            self.pdf_file.seek(initial_pos)
            return magic_number == b'%PDF-'
        except Exception:
            return False

    def extrair_tabelas(self):
        """
        Retorna BytesIO com Excel (openpyxl-saved) ou None em caso de falha.
        """
        try:
            if not self._validate_pdf_magic_number():
                return None

            # Camelot expects a filename; self.pdf_file.name should exist if file was opened from path
            pdf_path = getattr(self.pdf_file, "name", None)
            if not pdf_path or not os.path.exists(pdf_path):
                # fallback: dump to temp file and use that path
                self.pdf_file.seek(0)
                pdf_bytes = self.pdf_file.read()
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                    tmp.write(pdf_bytes)
                    tmp_path = tmp.name
                use_path = tmp_path
                remove_tmp = True
            else:
                use_path = pdf_path
                remove_tmp = False

            # Escreve as Tabelas
            tabelas = camelot.read_pdf(use_path, pages="all")
            if remove_tmp:
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass

            if len(tabelas) == 0:
                return None

            dfs = []
            new_column_titles = [
                '  ', 'Proposição', 'Autoria', 'Regime', 'Descrição',
                'Relator', 'Solicitações de Pauta', 'Assessor', 'Orientação', 'Observações'
            ]

            for i, tabela in enumerate(tabelas):
                df = tabela.df
                if i == 0:
                    dfs.append(df)
                else:
                    if len(df) > 1:
                        dfs.append(df.iloc[1:].reset_index(drop=True))
                    else:
                        dfs.append(df.reset_index(drop=True))

            df_final = pd.concat(dfs, ignore_index=True)

            if not df_final.empty:
                if not df_final.columns.empty:
                    first_row_str = df_final.iloc[0].astype(str).str.cat(sep=' ')
                    if any(title.lower() in first_row_str.lower() for title in new_column_titles):
                        df_final = df_final.iloc[1:].reset_index(drop=True)

                current_cols = df_final.columns.tolist()
                num_cols_to_assign = min(len(current_cols), len(new_column_titles))
                df_final.columns = new_column_titles[:num_cols_to_assign]

                if len(current_cols) < len(new_column_titles):
                    for j in range(len(current_cols), len(new_column_titles)):
                        df_final[new_column_titles[j]] = ''

                if all(col in df_final.columns for col in new_column_titles):
                    df_final = df_final[new_column_titles]
                else:
                    df_final = df_final[df_final.columns.intersection(new_column_titles)]
                    missing_cols = [col for col in new_column_titles if col not in df_final.columns]
                    for m_col in missing_cols:
                        df_final[m_col] = ''
                    df_final = df_final[new_column_titles]

            # Save to BytesIO via a temp file (openpyxl loads more reliably from path)
            temp_excel_path = "temp_excel_output.xlsx"
            df_final.to_excel(temp_excel_path, sheet_name="Tabelas", index=False, header=True)

            workbook = openpyxl.load_workbook(temp_excel_path)
            sheet = workbook["Tabelas"]
            font = Font(name='Arial', size=14)
            bold_font = Font(name='Arial', size=14, bold=True)
            light_grey_fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")
            sheet.insert_rows(idx=1,amount=1)
            sheet["A1"] = f"REUNIÃO DE LÍDERES - {data1}"

            for row in sheet.iter_rows():
                for cell in row:
                    if cell.row <= 2:
                        cell.font = bold_font
                        cell.fill = light_grey_fill
                        cell.alignment = Alignment(wrapText=True, horizontal='center', vertical='center')
                    else:
                        cell.font = font
                        cell.alignment = Alignment(wrapText=True, horizontal='left', vertical='top')

            if sheet.max_column >= 1:
                for cell in sheet['A']:
                    cell.fill = light_grey_fill

            for col in sheet.columns:
                max_length = 0
                for cell in col:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                col_letter = col[0].column_letter
                sheet.column_dimensions[col_letter].width = min(max_length * 1.2, 80)

            num_columns_in_excel = sheet.max_column
            if num_columns_in_excel >= 1:
                for i in range(1, 10):
                    col_index_from_end = num_columns_in_excel - i
                    col_letter = get_column_letter(col_index_from_end + 1)
                    sheet.column_dimensions[col_letter].width = 18  # Specific width for the last five columns
                    sheet.column_dimensions['E'].width = 70
                    sheet.column_dimensions['A'].width = 5
                    sheet.row_dimensions[1].height = 30

            sheet.merge_cells(start_row=1, start_column=1,end_row=1,end_column=sheet.max_column)
            output_final = BytesIO()
            workbook.save(output_final)
            output_final.seek(0)
            try:
                os.remove(temp_excel_path)
            except Exception:
                pass
            return output_final

        except Exception as e:
            print(f"[Erro extrair_tabelas] {e}")
            return None

    def extrair_blocos_por_numeros(self, stop_word=None):
        """
        Retorna lista de blocos (numero, texto) ou [] em caso de falha.
        """
        try:
            if not self._validate_pdf_magic_number():
                return []

            # read bytes from start
            self.pdf_file.seek(0)
            pdf_bytes = self.pdf_file.read()

            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                temp_pdf.write(pdf_bytes)
                temp_path = temp_pdf.name

            texto_total = ""
            with pdfplumber.open(temp_path) as pdf:
                for pagina in pdf.pages:
                    try:
                        texto = pagina.extract_text()
                    except Exception:
                        texto = None
                    if not texto:
                        continue
                    if stop_word and stop_word in texto:
                        break
                    texto_total += "\n" + texto

            try:
                os.remove(temp_path)
            except Exception:
                pass

            linhas = [linha.strip() for linha in texto_total.split("\n")]
            linhas = [linha for linha in linhas if linha or re.fullmatch(r"^\d+\.?$", linha.strip())]

            blocos = []
            buffer = []
            numero_atual = None
            started_block = False

            numero_com_texto_regex = re.compile(r"^(\d+)\.\s*(.*)")
            numero_linha_isolada_regex = re.compile(r"^\d+\.?$")

            for linha in linhas:
                linha_strip = linha.strip()
                match_com_texto = numero_com_texto_regex.match(linha_strip)
                match_linha_isolada = numero_linha_isolada_regex.match(linha_strip)

                if match_com_texto:
                    if started_block and numero_atual is not None:
                        blocos.append((numero_atual, " ".join(buffer).strip()))
                    buffer = [match_com_texto.group(2).strip()]
                    numero_atual = int(match_com_texto.group(1))
                    started_block = True

                elif match_linha_isolada:
                    if started_block and numero_atual is not None:
                        blocos.append((numero_atual, " ".join(buffer).strip()))
                    buffer = []
                    numero_atual = int(linha_strip.replace('.', ''))
                    started_block = True

                elif started_block:
                    buffer.append(linha_strip)

            if started_block and numero_atual is not None and buffer:
                blocos.append((numero_atual, " ".join(buffer).strip()))

            blocos = [(num, text) for num, text in blocos if text.strip()]
            return blocos

        except Exception as e:
            print(f"[ERRO] Falha ao extrair blocos numerados: {e}")
            return []

# -----------------------
# Document Generator
# -----------------------
class DocumentGenerator:

    def resource_path(relative_path):
        """ Pegando o caminho absoluto para o recurso."""
        try:
            # Cria uma pasta temporária e armazena o caminho em _MEIPASS
            base_path = sys._MEIPASS
        except Exception:
            base_path = os.path.abspath(".")
        return os.path.join(base_path, relative_path)

    def gerar_word_com_blocos(self, blocos):
        image_path1 = resource_path("images/Oposicao.png")
        image_path2 = resource_path("images/lideranca.png")

        try:
            doc = Document()
            section = doc.sections[0]
            original_width, original_height = section.page_width, section.page_height
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = original_height
            section.page_height = original_width
            header = section.header
            paragraph = header.add_paragraph()
            heading = paragraph.add_run()
            #heading.add_picture(image_path1, width=Pt(80), height=Pt(80)) # Inserindo a image no Cabeçalho
            heading = paragraph.add_run(f'      PAUTA DE PLENÁRIO - {data1}      ')
            #heading.add_picture(image_path2, width=Pt(80), height=Pt(80)) # Inserindo a image no Cabeçalho
            custom = heading.font
            custom.name = "Arial"
            custom.size = Pt(13)
            custom.bold = True
            custom.color.rgb = RGBColor(0x00, 0x00, 0xFF)

            tabela = doc.add_table(rows=1, cols=2)
            tabela.style = "Table Grid"
            hdr_cells = tabela.rows[0].cells
            hdr_cells[0].width = Inches(-6)
            hdr_cells[0].text = "Projeto"
            hdr_cells[1].text = "Análise"

            texto_complementa = "Autoria: \nRelatoria: \nAssessoria Oposição: \nMinoria: \nPOSICIONAMENTO:"

            for cell in hdr_cells:
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:fill"), "D3D3D3")
                cell._tc.get_or_add_tcPr().append(shading_elm)
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(13)

            for numero, texto in blocos:
                row_cells = tabela.add_row().cells
                row_cells[0].text = f"{numero}. {texto}\n {texto_complementa}"
                for c in row_cells:
                    for run in c.paragraphs[0].runs:
                        run.font.size = Pt(10)
                        run.font.name = "Arial"

            output = BytesIO()
            doc.save(output)
            output.seek(0)
            return output
        except Exception as e:
            print(f"Erro ao gerar documento Word: {e}")
            return None

# -----------------------
# Worker Thread
# -----------------------
class ProcessWorker(QThread):
    finished_signal = pyqtSignal(object, str, str)  # BytesIO, suggested_name, type
    error_signal = pyqtSignal(str)

    def __init__(self, pdf_path, output_type):
        super().__init__()
        self.pdf_path = pdf_path
        self.output_type = output_type  # 'excel' or 'word'

    def run(self):
        try:
            # Open file in worker thread (independent file object)
            with open(self.pdf_path, "rb") as f:
                extractor = PDFExtractor(f)
                # validate
                if not extractor._validate_pdf_magic_number():
                    self.error_signal.emit("Arquivo inválido (não parece ser um PDF).")
                    return

                if self.output_type == "excel":
                    excel_io = extractor.extrair_tabelas()
                    if not excel_io:
                        self.error_signal.emit("Nenhuma tabela encontrada ou falha na extração.")
                        return
                    suggested = f"{data} - REUNIÃO DE LÍDERES.xlsx"
                    self.finished_signal.emit(excel_io, suggested, "excel")
                else:
                    blocos = extractor.extrair_blocos_por_numeros(stop_word="AVISO")
                    if not blocos:
                        self.error_signal.emit("Nenhum bloco numerado encontrado.")
                        return
                    gen = DocumentGenerator()
                    word_io = gen.gerar_word_com_blocos(blocos)
                    if not word_io:
                        self.error_signal.emit("Falha ao gerar o documento Word.")
                        return
                    suggested = f"{data} - Pauta plenário.docx"
                    self.finished_signal.emit(word_io, suggested, "word")
        except Exception as e:
            tb = traceback.format_exc()
            self.error_signal.emit(f"Erro durante processamento: {e}\n{tb}")


# -----------------------
# Main App (GUI)
# -----------------------
class PDFApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle(" Extrator de PDFs")
        self.setGeometry(300, 200, 700, 420)

        self._timer = None
        self.worker = None

        layout = QVBoxLayout()

        title = QLabel("📂 Extrator de Dados de PDF")
        title.setFont(QFont("Arial", 16, QFont.Bold))
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)

        # input editable
        self.file_input = QLineEdit()
        self.file_input.setPlaceholderText("Digite ou cole o caminho do PDF ou clique em 'Procurar'...")
        self.file_input.setStyleSheet("font: 12pt Arial;")
        layout.addWidget(self.file_input)

        file_btn = QPushButton("📎 Procurar PDF")
        file_btn.setStyleSheet("background-color: #2980b9; color: white; padding: 8px; font-size: 16px;")
        file_btn.clicked.connect(self.selecionar_pdf)
        layout.addWidget(file_btn)

        group_box = QGroupBox("Formato de Saída")
        group_box.setStyleSheet("font-size: 14px;")
        group_layout = QHBoxLayout()
        self.option_excel = QRadioButton("Excel")
        self.option_word = QRadioButton("Word")
        self.option_excel.setStyleSheet("font-size: 14px;")
        self.option_word.setStyleSheet("font-size: 14px;")
        self.option_excel.setChecked(True)
        group_layout.addWidget(self.option_excel)
        group_layout.addWidget(self.option_word)
        group_box.setLayout(group_layout)
        layout.addWidget(group_box)

        self.progress = QProgressBar()
        self.progress.setValue(0)
        layout.addWidget(self.progress)

        process_btn = QPushButton("▶ Processar PDF")
        process_btn.setStyleSheet("background-color: #27ae60; color: white; padding: 10px; font-size: 16px;")
        process_btn.clicked.connect(self.processar_pdf)
        layout.addWidget(process_btn)

        self.status = QLabel("")
        self.status.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.status)


        # small footer with log button optionally
        self.log_btn = QPushButton("Ver logs (abrir app_log.txt)")
        #self.log_btn.clicked.connect(self.open_log)
        #self.log_btn.setStyleSheet("background-color: #7f8c8d; color: white; padding: 6px;")
        #layout.addWidget(self.log_btn)


        self.setLayout(layout)

    # -----------------------
    # UI helpers
    # -----------------------

    def selecionar_pdf(self):
        pdf_path, _ = QFileDialog.getOpenFileName(self, "Selecionar PDF", "", "PDF Files (*.pdf)")
        if pdf_path:
            self.file_input.setText(pdf_path)


    def get_pdf_path(self):
        return self.file_input.text().strip()


    def set_ui_enabled(self, enabled: bool):
        # Habilita/desabilita controles para evitar reentrada durante o processamento
        self.file_input.setEnabled(enabled)
        for w in self.findChildren(QPushButton):
            if w is not self.log_btn:
                w.setEnabled(enabled)
        for r in (self.option_excel, self.option_word):
            r.setEnabled(enabled)

    def open_log(self):
        # Abre o arquivo de log com o editor padrão, se existir, caso contrário, mostre a mensagem
        log_path = os.path.abspath("app_log.txt")
        if not os.path.exists(log_path):
            QMessageBox.information(self, "Logs", "Ainda não existem logs.")
            return
        try:
            if sys.platform.startswith("win"):
                os.startfile(log_path)
            elif sys.platform == "darwin":
                os.system(f"open '{log_path}'")
            else:
                os.system(f"xdg-open '{log_path}'")
        except Exception as e:
            QMessageBox.warning(self, "Erro", f"Não foi possível abrir o log: {e}")

    def registrar_log(self, mensagem):
        pass
        '''
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        log_entry = f"[{timestamp}] {mensagem}\n"
        try:
            with open("app_log.txt", "a", encoding="utf-8") as log_file:
                log_file.write(log_entry)
        except Exception as e:
            print(f"⚠ Erro ao gravar log: {e}")
        '''
    # -----------------------
    # Animação da barra de progresso
    # -----------------------
    def animate_progress(self, start, end, step=3, delay=30, auto_reset=False, reset_delay=2000):
        self.progress.setValue(start)

        def update():
            value = self.progress.value() + step if end > start else self.progress.value() - step
            if (end > start and value >= end) or (end < start and value <= end):
                self.progress.setValue(end)
                timer.stop()

                if end == 100:
                    self.status.setText("✅ Processo concluído com sucesso!")
                    self.status.setStyleSheet("font-size: 14px;")
                    if auto_reset:
                        QTimer.singleShot(reset_delay, self.reset_progress_after_delay)
            else:
                self.progress.setValue(value)

        timer = QTimer(self)
        timer.timeout.connect(update)
        timer.start(delay)

    def reset_progress_after_delay(self):
        def update_back():
            value = self.progress.value() - 3
            if value <= 0:
                self.progress.setValue(0)
                timer.stop()
                self.status.setText("")  # limpa mensagem ao resetar
            else:
                self.progress.setValue(value)

        timer = QTimer(self)
        timer.timeout.connect(update_back)
        timer.start(30)

    # -----------------------
    # Worker callbacks
    # -----------------------
    def on_worker_finished(self, bytes_io, suggested_name, out_type):
        try:
            # stop any animation timer
            if self._timer and self._timer.isActive():
                self._timer.stop()

            # animar até quase completar (visual)
            current = self.progress.value()
            target_mid = max(current, 80)
            self.animate_progress(current, target_mid)

            # Pergunta como que salvar (main thread GUI)
            if out_type == "excel":
                filter_str = "Arquivos Excel (*.xlsx)"
                default = suggested_name
            else:
                filter_str = "Arquivos Word (*.docx)"
                default = suggested_name

            output_name, _ = QFileDialog.getSaveFileName(self, "Salvar Arquivo", default, filter_str)

            if not output_name:
                # user canceled save
                self.animate_progress(target_mid, 0)
                self.status.setText("⚠ Operação cancelada pelo usuário")
                #self.registrar_log("Usuário cancelou a exportação (salvar dialog).")
                self.set_ui_enabled(True)
                return


            try:
                with open(output_name, "wb") as f:
                    f.write(bytes_io.getvalue())
            except Exception as e:
                self.animate_progress(target_mid, 0)
                self.status.setText("❌ Falha ao salvar o arquivo")
                #self.registrar_log(f"Erro ao salvar arquivo: {e}\n{traceback.format_exc()}")
                QMessageBox.critical(self, "Erro", f"Não foi possível salvar o arquivo:\n{e}")
                self.set_ui_enabled(True)
                return

            # success
            self.animate_progress(target_mid, 100)
            #self.status.setText(f"✅ Arquivo salvo em {output_name}")
            #self.registrar_log(f"Arquivo salvo em {output_name}")
            self.animate_progress(20, 100, auto_reset=True)
        finally:
            # Reativar a interface do usuário e o trabalhador de limpeza
            self.set_ui_enabled(True)
            try:
                if self.worker:
                    self.worker.quit()
                    self.worker.wait(100)
            except Exception:
                pass
            self.worker = None

    def on_worker_error(self, message):
        # Para Animação
        try:
            if self._timer and self._timer.isActive():
                self._timer.stop()
        except Exception:
            pass

        self.animate_progress(self.progress.value(), 0)
        #self.status.setText("❌ Erro — veja o log")
        #self.registrar_log(f"Worker error: {message}")
        QMessageBox.critical(self, "Erro", message)
        self.set_ui_enabled(True)
        # Limpando a Barra
        try:
            if self.worker:
                self.worker.quit()
                self.worker.wait(100)
        except Exception:
            pass
        self.worker = None

    # -----------------------
    # Main action
    # -----------------------
    def processar_pdf(self):
        pdf_path = self.get_pdf_path()
        if not pdf_path or not os.path.exists(pdf_path):
            QMessageBox.warning(self, "Erro", "Digite ou selecione um PDF válido!")
            #self.registrar_log("Arquivo PDF inválido ou não encontrado.")
            return

        # Desativar a UI durante a execução
        self.set_ui_enabled(False)
        self.status.setText("🔄 Iniciando processamento...")
        self.status.setStyleSheet("font-size: 14px;")
        self.animate_progress(0, 20)

        out_type = "excel" if self.option_excel.isChecked() else "word"

        # Create and start worker
        self.worker = ProcessWorker(pdf_path, out_type)
        self.worker.finished_signal.connect(self.on_worker_finished)
        self.worker.error_signal.connect(self.on_worker_error)
        self.worker.start()

# -----------------------
# Run
# -----------------------
if __name__ == "__main__":
    # === Define um AppUserModelID único para o Windows reconhecer o ícone ===
    try:
        myappid = "com.extrator.pdf"  # pode ser qualquer string única
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
    except Exception as e:
        print("Aviso: não foi possível definir AppUserModelID:", e)

    # === Função para localizar arquivos mesmo no modo .exe ===
    def resource_path(relative_path):
        """
        Retorna o caminho absoluto de um recurso,
        compatível tanto com o script Python quanto com o .exe do PyInstaller.
        """
        if hasattr(sys, "_MEIPASS"):
            # PyInstaller armazena recursos temporariamente nessa pasta
            return os.path.join(sys._MEIPASS, relative_path)
        return os.path.join(os.path.abspath("."), relative_path)

    # === Inicializa o aplicativo ===
    app = QApplication(sys.argv)
    window = PDFApp()

    # === Define o ícone do app (janela + taskbar) ===
    icon_path = resource_path("images/icon.png")
    app.setWindowIcon(QIcon(icon_path))
    window.setWindowIcon(QIcon(icon_path))

    # === Exibe a janela ===
    window.show()

    # === Executa o loop principal ===
    sys.exit(app.exec_())